# -*- coding: utf-8 -*-
"""生成 CRM 企业级智能销售平台 Word 报告文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ====== 样式设置 ======
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
# Set East Asian font
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = doc.styles['Normal'].element.get_or_add_rPr().makeelement(qn('w:rFonts'), {})
    rPr.insert(0, rFonts)
rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = run._element.get_or_add_rPr().makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    if align is not None: p.alignment = align
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def screenshot_hint(caption):
    p = doc.add_paragraph()
    run = p.add_run(f'📷 【截图位：{caption}】')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
add_para('CRM 企业级智能销售平台', bold=True, size=28, color=(37, 99, 235), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('系统开发完整报告 v3.0', bold=True, size=16, color=(100, 100, 100), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('包含：CRM管理后台 + 用户会员门户 + 全数据库 + AI智能体', size=12, color=(100, 100, 100), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('报告日期：2026年5月31日', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('技术栈：Node.js + Express + SQLite + Vue 3 + Element Plus', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ============================================================
# 目录
# ============================================================
add_heading('目录', 1)
toc_items = [
    '一、系统概述', '二、技术架构', '三、CRM管理后台功能详解',
    '四、用户会员门户功能详解', '五、AI智能体功能', '六、数据库设计',
    '七、API端点清单', '八、UI设计体系', '九、启动与部署',
    '十、项目管理文件结构', '十一、数据清单', '十二、开发统计'
]
for t in toc_items:
    add_para(t, size=12)
doc.add_page_break()

# ============================================================
# 一、系统概述
# ============================================================
add_heading('一、系统概述', 1)

add_heading('1.1 系统定位', 2)
add_para('CRM企业级智能销售平台是一套面向大学生零售场景的全栈客户关系管理系统，定位为"管理后台 + 会员门户"双端一体化解决方案。系统涵盖客户全生命周期管理（获客→转化→成交→服务→分析），并集成了AI智能助手、消遣模块等创新功能。')

add_heading('1.2 当前运行服务', 2)
add_table(
    ['服务', '端口', '技术栈', '登录账号', '密码'],
    [
        ['CRM 管理后台', '3002', 'Vue 3 + Element Plus', 'admin', 'admin123'],
        ['用户会员门户', '3003', 'Vue 3 + Element Plus', '13800000000', '123456'],
        ['后端 API', '8081', 'Node.js + Express', '-', '-'],
    ],
    [3, 2, 4, 3, 3]
)
screenshot_hint('浏览器同时打开三个服务的效果截图')

add_heading('1.3 业务闭环', 2)
add_para('线索获取 → 商机转化 → 报价谈判 → 订单成交 → 客户服务 → 数据分析 → 营销触达 → 复购转化')
add_para('以上闭环完整覆盖了CRM系统的核心业务流程，CRM管理后台负责B端运营管理，会员门户负责C端用户消费体验。')

doc.add_page_break()

# ============================================================
# 二、技术架构
# ============================================================
add_heading('二、技术架构', 1)

add_heading('2.1 架构图', 2)
add_para('┌────────────────────────────────────────────────────┐')
add_para('│          浏览器访问层                              │')
add_para('│  CRM后台(:3002)  │  会员门户(:3003)               │')
add_para('└────────┬────────────────────┬──────────────────────┘')
add_para('         │  HTTP REST API    │')
add_para('┌────────▼────────────────────▼──────────────────────┐')
add_para('│       Express API Server (port 8081)               │')
add_para('│       JWT认证 / 路由分发 / 业务逻辑                │')
add_para('└────────────────────────┬───────────────────────────┘')
add_para('                         │ better-sqlite3')
add_para('┌────────────────────────▼───────────────────────────┐')
add_para('│           SQLite 数据库 (data/crm.db)              │')
add_para('│           75+ 张表 / 持久化存储                    │')
add_para('└────────────────────────────────────────────────────┘')

add_heading('2.2 技术选型', 2)
add_table(
    ['层级', '技术', '版本', '说明'],
    [
        ['后端运行时', 'Node.js', '22.22.2', '高性能异步I/O'],
        ['后端框架', 'Express', '4.x', '轻量级路由框架'],
        ['数据库', 'SQLite (better-sqlite3)', '11.x', '零配置本地数据库'],
        ['认证', 'JWT (jsonwebtoken)', '9.x', '24h有效期Token'],
        ['密码加密', 'bcrypt', '5.x', '单向哈希加密'],
        ['管理前端', 'Vue 3', '3.x', 'Composition API'],
        ['UI库', 'Element Plus', '2.x', '企业级组件库'],
        ['图表', 'ECharts', '6.x', 'RFM/漏斗图表'],
        ['构建工具', 'Vite', '8.x', '极速构建'],
        ['会员前端', 'Vue 3 + Element Plus', '-', '移动端适配'],
    ],
    [3, 4, 2, 5]
)

doc.add_page_break()

# ============================================================
# 三、CRM管理后台功能详解
# ============================================================
add_heading('三、CRM管理后台功能详解', 1)
add_para('CRM管理后台是B端运营管理核心，提供29个功能模块，覆盖从客户获取到数据分析的全链路。登录后进入深蓝色侧边栏布局，系统设置区包含门户管理、积分管理、优惠券管理等特色模块。')

screenshot_hint('CRM后台登录后的仪表盘首页截图')

# 3.1 工作台
add_heading('3.1 工作台 — 仪表盘', 2)
add_para('仪表盘是系统的数据总览中心，顶部展示6个核心指标卡片（客户总数、线索数、商机数、订单数、月销售额、待办任务），下方为可拖拽的自定义小部件区域。集成了消遣功能（捏泡泡、番茄钟）。数据实时从后端统计。')
add_table(
    ['指标', '当前值', '说明'],
    [['客户总数', '12', '含高价值/普通/沉睡各等级'],
     ['线索数', '5', '待转化商机的潜在客户'],
     ['商机数', '5', '处于不同销售阶段'],
     ['订单总数', '56', '含已付款/已发货/已完成'],
     ['月销售额', '动态', 'SUM(total_amount)'],
     ['待办任务', '2', '跟进线索+月度报告']],
    [3, 2, 7]
)
screenshot_hint('仪表盘页面全貌')

# 3.2 客户管理
add_heading('3.2 客户管理', 2)
add_para('客户管理支持卡片视图和表格视图双模式切换。卡片模式显示客户头像首字母（渐变色背景）、姓名、电话、等级标签、行业标签，悬停时卡片上浮并显示蓝色渐变顶部条。表格模式提供标准的数据表格视图，支持搜索、筛选、分页。')
add_para('关键功能：')
add_para('• 客户信息完整字段：姓名、电话、邮箱、性别、年龄、来源、等级、行业、地址、备注等')
add_para('• 智能标签系统：8种标签（高频消费、美妆爱好者、文具达人、价格敏感、新生、毕业生等）')
add_para('• 学生画像：年级、专业类别、生活费区间、校区、宿舍信息')
add_para('• RFM客户分层：核心价值客户/潜力客户/一般客户/流失风险客户')
add_para('• 客户共享与保护：支持客户在销售人员间共享，设置保护期')
screenshot_hint('客户管理卡片视图')

# 3.3 销售管理
add_heading('3.3 销售管理', 2)
add_para('销售管理包含线索管理、商机管理、报价管理、订单管理、跟进记录五大子模块，覆盖完整的销售漏斗。')

add_heading('3.3.1 线索管理', 3)
add_para('线索是潜在客户入口。支持"我的线索"和"公海池"双Tab切换。每条线索包含名称、来源（转介/展会/网站/地推/活动）、联系人信息、跟进次数、下次跟进时间。可一键转化为商机。')
add_table(
    ['线索名称', '来源', '联系人', '跟进次数', '状态'],
    [['陈思雨-美妆批发', '转介', '陈思雨', '2', '跟进中'],
     ['刘建国-企业采购', '展会', '刘建国', '3', '跟进中'],
     ['新锐科技-办公用品', '网站', '赵总', '1', '跟进中'],
     ['星光百货-合作', '地推', '李总', '1', '跟进中'],
     ['秦岚-行政采购', '门店', '秦岚', '1', '跟进中']],
    [3, 2, 2, 2, 2]
)
screenshot_hint('线索管理页面')

add_heading('3.3.2 商机管理', 3)
add_para('商机管理支持看板视图和列表视图。看板按销售阶段分列（初步接触→需求确认→方案报价→商务谈判→合同签约），支持拖拽改变阶段。每个商机记录名称、关联客户、预计金额、成交概率、预计成交日期。')
add_table(
    ['商机名称', '客户', '金额', '阶段', '概率'],
    [['陈思雨年度美妆采购', '陈思雨', '¥50,000', '需求确认', '60%'],
     ['刘建国企业办公用品年框', '刘建国', '¥120,000', '方案报价', '40%'],
     ['李佳琪美妆品牌合作', '李佳琪', '¥200,000', '商务谈判', '70%'],
     ['蒋晓晓季节新品采购', '蒋晓晓', '¥80,000', '需求确认', '50%'],
     ['秦岚公司季度采购', '秦岚', '¥60,000', '方案报价', '45%']],
    [3, 2, 2, 2, 2]
)
screenshot_hint('商机看板视图')

add_heading('3.3.3 报价管理', 3)
add_para('报价管理提供完整的报价生命周期管理：草稿→已提交→已批准/已驳回→转订单。报价单包含客户信息、商品明细（产品名称、规格、数量、单价、折扣）、合计金额、有效期。支持一键审批流转和转换为正式销售订单。')
screenshot_hint('报价管理页面')

add_heading('3.3.4 订单管理', 3)
add_para('当前系统共有56笔订单，覆盖已付款、已发货、已完成等多状态。每笔订单记录客户、金额、状态、创建时间。订单详情展示商品明细（产品名称×数量×单价）。支持按状态筛选、取消订单、确认收货、申请退换货。')
add_table(
    ['状态', '说明', '数量'],
    [['已付款', '已支付待发货', '约30笔'],
     ['已发货', '已发货待收货', '约10笔'],
     ['已完成', '交易完成', '约10笔'],
     ['已取消', '已取消订单', '少量'],
     ['退换货中', '申请退换处理中', '少量']],
    [3, 4, 4]
)
screenshot_hint('订单管理列表')

add_heading('3.3.5 跟进记录', 3)
add_para('跟进记录记录每次与客户的互动，包括拜访、回访、电话沟通等类型。支持记录跟进内容、计划时间、实际时间，帮助销售团队有效管理客户关系。')

# 3.4 市场运营
add_heading('3.4 市场运营', 2)
add_para('市场运营模块包含市场活动、营销活动、优惠券三个子模块。')

add_heading('3.4.1 市场活动', 3)
add_para('已预置4场市场活动：')
add_table(
    ['活动名称', '类型', '预算', '时间', '状态'],
    [['校园春季推广', '校园推广', '¥5,000', '2026.4-5', '已完成'],
     ['会员日满减大促', '促销', '¥8,000', '2026.5-6', '进行中'],
     ['新品体验官招募', '新品', '¥3,000', '2026.6', '计划中'],
     ['毕业季感恩回馈', '促销', '¥6,000', '2026.6-7', '计划中']],
    [3, 2, 2, 3, 2]
)
screenshot_hint('市场活动列表')

add_heading('3.4.2 营销活动', 3)
add_para('支持满减、折扣、限时等多种营销方式。可创建优惠券并发放给客户。')
add_table(
    ['活动名称', '类型', '规则', '状态'],
    [['五一狂欢节满减', '满减', '全场满200减50', '进行中'],
     ['高价值客户专属福利', '促销', '高价值客户专享8折', '进行中'],
     ['初夏新品首发', '新品', '新品上市限时优惠', '进行中']],
    [3, 2, 4, 2]
)

add_heading('3.4.3 优惠券管理', 3)
add_para('系统设置中新增优惠券管理模块（见3.8.3节），支持优惠券模板CRUD、一键发放到客户、发放记录追踪。')

# 3.5 协同办公
add_heading('3.5 协同办公', 2)
add_para('协同办公模块包含日程任务、公告通知、站内信、文档管理、审批管理五个子模块。')
add_table(
    ['模块', '数据量', '功能说明'],
    [['日程任务', '7+2条', '客户回访日程、待办任务管理'],
     ['公告通知', '2条', '618大促安排、新品上线通知'],
     ['站内信', '1封', '内部邮件系统'],
     ['文档管理', '-', '文件上传与管理'],
     ['审批管理', '-', '审批流配置与记录']],
    [3, 2, 8]
)
screenshot_hint('协同办公-日程任务页面')

# 3.6 数据分析
add_heading('3.6 数据分析', 2)
add_para('数据分析是系统的核心决策支持模块，包含RFM分析、销售预测、业绩排行、学生分析四个子模块。')

add_heading('3.6.1 RFM分析', 3)
add_para('RFM分析基于R（最近消费间隔）、F（消费频率）、M（消费金额）三维度对客户进行分层评估。页面包含：')
add_para('• 顶部统计卡片：4级分层（核心价值客户/潜力客户/一般客户/流失风险客户）各多少人')
add_para('• 左侧散点图：X轴=R值，Y轴=F值，气泡大小=M值，不同颜色代表不同分层')
add_para('• 右侧饼图：客户分层分布百分比')
add_para('• 底部明细表：每位客户的R/F/M得分、综合分层、等级、来源、计算时间')
add_para('• 支持按分层筛选、数据自动从crm_rfm_score表读取（关联customer表）')
screenshot_hint('RFM分析完整页面（散点图+饼图+表格）')

add_heading('3.6.2 销售预测', 3)
add_para('展示12个月历史销售趋势折线图 + 3个月预测趋势，支持切换折线/柱状图。数据来源于sales_order表的历史订单统计。')
screenshot_hint('销售预测趋势图')

add_heading('3.6.3 业绩排行', 3)
add_para('按月份展示各员工的销售业绩排行，包含销售额、新客户数等指标。')
screenshot_hint('员工业绩排行')

add_heading('3.6.4 学生分析 🎓', 3)
add_para('学生分析是面向大学生客户的专项分析仪表盘，包含：')
add_para('• 概览卡片：学生客户数、已认证数、RFM分层数、标签总数')
add_para('• RFM客户分层表格：可点击"重新计算RFM"按钮')
add_para('• 学生画像表：展示年级、专业、生活费、校区、宿舍')
add_para('• 校园活动管理：可新建校园推广活动')
add_para('• 营销模板库：满减、半价、闪购模板')
add_para('• 自动化工作流：新客欢迎/沉睡唤醒/生日关怀')
add_para('• 校区商圈分析：主校区5000人/大学城12000人')
screenshot_hint('学生分析仪表盘全貌')

# 3.7 服务管理
add_heading('3.7 服务管理', 2)
add_para('服务管理模块处理客户服务请求和工单。当前有2条服务请求（产品使用咨询、退换货处理）。支持工单分类（咨询/退换货/质量问题）、优先级设置、处理状态跟踪。')
screenshot_hint('服务管理页面')

# 3.8 系统设置 — 特色管理模块
add_heading('3.8 系统设置 — 特色管理模块', 2)

add_heading('3.8.1 🛒 门户管理', 3)
add_para('门户管理是CRM后台控制会员门户的核心模块，功能包括：')
add_para('• 概览面板：注册用户数、上架商品数、累计营收、今日签到')
add_para('• 首页Banner编辑：可添加/删除/编辑轮播图（标题+副标题+颜色），保存后会员门户首页实时生效')
add_para('• 首页公告编辑：编辑顶部公告文字，实时更新')
add_para('• 产品上下架管理：以表格形式列出28款产品，可直接切换上架/下架开关，设置会员价')
add_para('• 注册用户管理：查看所有门户注册用户（昵称、手机号、积分、消费、注册时间）')
screenshot_hint('门户管理页面：Banner编辑+产品上下架表格')

add_heading('3.8.2 💰 积分管理', 3)
add_para('积分管理提供完整的积分运营功能：')
add_para('• 概览面板：当前积分池 / 累计发放 / 累计消耗 / 今日签到')
add_para('• 积分明细表：200条历史积分记录（客户名、类型、积分、原因、来源）')
add_para('• 筛选功能：支持按客户名称远程搜索过滤')
add_para('• 手动调整：弹窗选择客户 → 选择增减 → 输入积分数量 → 填写原因 → 保存（同步更新门户用户积分）')
screenshot_hint('积分管理页面：概览+明细表+手动调整弹窗')

add_heading('3.8.3 🎫 优惠券管理', 3)
add_para('优惠券管理提供完整的优惠券生命周期管理：')
add_para('• 概览面板：券模板数 / 已发放 / 已使用 / 未使用')
add_para('• 券模板Tab：创建/编辑/删除优惠券（名称、面额、使用门槛、发行数量）')
add_para('• 发放记录Tab：查看所有发放记录（券名、客户、状态、时间）')
add_para('• 一键发放：选择券模板 → 远程搜索客户 → 确认发放（同时写入coupon_record和portal_coupon_grant两张表，实现双系统数据同步）')
screenshot_hint('优惠券管理：模板列表+发放记录+发放弹窗')

add_heading('3.8.4 💬 反馈管理', 3)
add_para('反馈管理处理会员门户用户提交的反馈：')
add_para('• 概览统计：全部/未处理/处理中/已回复四分类计数，点击切换')
add_para('• 状态流转：未处理 → 标记处理中')
add_para('• 回复功能：弹窗输入回复内容 → 保存 → 自动变为"已回复"状态')
add_para('• 会员端同步：回复后会员在门户中立即看到商家回复（绿色背景卡片）')

add_heading('3.8.5 产品管理', 3)
add_para('产品管理维护28款商品（原8款基础产品 + 20款名创优品风格产品）。支持产品编码、名称、型号、分类、单位、售价、成本价、规格、描述等字段的CRUD操作。产品数据共享给会员门户使用。')

doc.add_page_break()

# ============================================================
# 四、用户会员门户功能详解
# ============================================================
add_heading('四、用户会员门户功能详解', 1)
add_para('用户会员门户是面向C端注册用户/付费会员的独立前端应用（端口3003），与CRM管理后台共享同一数据库。核心目标：提供会员注册登录、产品浏览、积分管理、会员权益展示、订单追踪等完整C端体验。')

add_heading('4.1 首页', 2)
add_para('首页采用高端消费风格设计，包含以下区块：')
add_para('• Hero Banner：橙色渐变背景 + "新品上市·限时特惠"主标题 + "立即抢购"按钮')
add_para('• ⚡ 限时秒杀区：深色半透明玻璃面板，6款商品半价，实时倒计时')
add_para('• 🎯 商品分类：胶囊标签式分类筛选（全部/家居/美妆/文具/数码/服饰）')
add_para('• 🔥 热销爆款：4列产品网格，每卡片含emoji图标、产品名、价格、已售数量')
add_para('• 🎖 会员权益：4级会员卡片（青铜🥉/白银🥈/黄金🥇/钻石💎），展示升级条件和权益')
screenshot_hint('会员门户首页完整截图')

add_heading('4.2 产品浏览', 2)
add_para('产品列表页支持：骨架屏加载动画、关键词搜索（400ms防抖）、分类胶囊筛选、4种排序（综合/价格升/价格降/销量）、会员价/原价双显示、秒杀/热卖标签。详情页采用双栏布局（左大图+右信息），含闪购价标签、销量统计、加入购物车/收藏按钮、星级评价、发表评价（+10积分奖励）。')
screenshot_hint('产品列表页和产品详情页')

add_heading('4.3 购物车与下单', 2)
add_para('购物车展示商品列表（emoji图标+名称+单价+数量调节+小计）。支持增量/减量、删除。底部合计栏显示"合计 ¥xxx 立即结算"。下单流程：购物车结算→自动创建sales_order+order_item记录→清空购物车→获得消费积分→跳转订单列表。')
screenshot_hint('购物车页面+下单成功提示')

add_heading('4.4 会员中心', 2)
add_para('会员中心展示当前等级大卡片（等级图标+权益数）+ 升级进度条（距离下一级还需消费xx元）。4级会员对比：')
add_table(
    ['等级', '名称', '升级条件', '折扣', '核心权益'],
    [['Lv.1', '青铜会员', '注册即享', '95折', '签到积分'],
     ['Lv.2', '白银会员', '累计≥500元', '92折', '双倍积分+生日礼包'],
     ['Lv.3', '黄金会员', '累计≥2000元', '88折', '专属客服'],
     ['Lv.4', '钻石会员', '累计≥5000元', '82折', '免运费+优先发货']],
    [2, 2, 3, 2, 5]
)
screenshot_hint('会员中心页面')

add_heading('4.5 积分中心', 2)
add_para('积分中心是会员活跃度的核心驱动模块：')
add_para('• 顶部紫色渐变Banner：显示可用积分和累计积分')
add_para('• 7天签到日历：已签到日✓绿色，今日待签蓝色边框，连续签到7天额外+20分')
add_para('• 🎁 积分商城：6款兑换商品（500-5000积分），点击确认兑换')
add_para('• 📋 积分明细：逐条显示积分变动记录（类型/分数/时间）')
add_table(
    ['兑换商品', '所需积分', '图标'],
    [['5元无门槛券', '500', '🎫'],
     ['10元满减券', '800', '💰'],
     ['品牌帆布袋', '1200', '🛍️'],
     ['限量徽章套装', '2000', '🏅'],
     ['保温杯(马卡龙)', '3000', '🥤'],
     ['蓝牙音箱', '5000', '🔊']],
    [3, 2, 2]
)
screenshot_hint('积分中心页面：签到+兑换+明细')

add_heading('4.6 领券中心', 2)
add_para('双Tab布局：可领取（展示所有可用优惠券 + 一键领取按钮）+ 我的优惠券（未使用/已使用状态标签）。')

add_heading('4.7 个人中心与辅助功能', 2)
add_para('个人中心聚合所有功能入口（订单/收藏/地址/优惠券/积分/会员/反馈），显示头像首字母、昵称、积分和消费统计。地址管理支持多地址CRUD、设默认。收藏夹展示已收藏商品列表。意见反馈支持提交（类型+文字+联系方式）+ 查看历史回复。')

doc.add_page_break()

# ============================================================
# 五、AI智能体功能
# ============================================================
add_heading('五、AI智能体功能', 1)
add_para('AI智能助手是CRM后台右下角的浮动对话机器人，支持15种自然语言指令。基于规则引擎，返回结构化数据。')

add_table(
    ['指令示例', '功能', '返回数据'],
    [['今日/今天', '工作概览', '日程数+待办数+客户数+线索数'],
     ['跟进/优先', '待跟进客户', '高价值客户列表'],
     ['预测/销售', '销售预测', '订单总数+总额+下月预测'],
     ['RFM/细分', 'RFM分析', '客户分层统计+4级人数'],
     ['统计/仪表盘', '数据总览', '客户/订单/线索/商机/产品全统计'],
     ['客户+姓名', '客户搜索', '匹配客户信息'],
     ['产品/商品', '产品查询', '产品总数+热销TOP3'],
     ['订单/成交', '订单概况', '总数+待付款数'],
     ['优惠券', '券统计', '模板数+发放数'],
     ['会员/VIP', '会员概况', '注册用户数+CRM客户数'],
     ['建议/优化', '业务建议', '5条基于数据分析的建议'],
     ['帮助/功能', '帮助列表', '全部15种指令清单'],
     ['你好/Hi', '问候', '欢迎语+引导']],
    [3, 3, 7]
)
screenshot_hint('AI智能助手对话窗口')

add_para('此外，系统还提供了以下AI辅助功能：')
add_para('• 销售预测：GET /api/ai/sales-forecast — 返回历史+预测趋势')
add_para('• 客户流失预警：GET /api/churn-predictions — 基于消费间隔评分（高/中/低风险）')
add_para('• 商品推荐：GET /api/ai-recommendations/:customerId — 基于随机+品类推荐4款商品')
add_para('• 营销文案生成：GET /api/ai-copy-templates — 预置微信/抖音/小红书风格文案')
add_para('• 销售话术：GET /api/ai/sales-script — 开场白/核心话术/异议处理/促成')

doc.add_page_break()

# ============================================================
# 六、数据库设计
# ============================================================
add_heading('六、数据库设计', 1)
add_para('系统使用SQLite数据库（文件位置：crm-server/data/crm.db），共75+张表，按业务域分类如下：')

add_table(
    ['类别', '表数', '核心表名'],
    [['CRM核心', '15', 'customer, sales_order, order_item, crm_lead, crm_lead_follow_up, crm_opportunity, crm_opportunity_team, crm_product, crm_quotation, crm_quotation_item, todo_task, crm_contact, crm_follow_up, crm_schedule, crm_task'],
     ['营销活动', '6', 'marketing_campaign, coupon, coupon_record, crm_marketing_activity, crm_marketing_template, crm_campus_campaign'],
     ['协同办公', '6', 'crm_notice, crm_notice_read, crm_internal_mail, crm_document, crm_approval_flow, crm_approval_record'],
     ['服务系统', '4', 'crm_service_request, crm_service_ticket, crm_service_ticket_ext, crm_return_request'],
     ['系统', '5', 'sys_user, sys_config, sys_dict_type, sys_dict_data, crm_sales_team, crm_sales_team_member'],
     ['大学生扩展', '20', 'crm_student_profile, crm_customer_tag, crm_rfm_score, crm_member_tier, crm_campus_cluster, crm_churn_prediction, crm_ai_recommendation, crm_ai_copy_template, crm_auto_workflow, crm_customer_care...'],
     ['会员门户', '12', 'portal_user, portal_user_auth, portal_address, portal_shopping_cart, portal_user_favorite, portal_browse_history, portal_product_review, portal_sign_record, portal_coupon_grant, portal_points_exchange, portal_feedback, portal_system_config'],
     ['其他', '8', 'crm_delivery, crm_delivery_item, crm_invoice, crm_receivable_plan, crm_receivable_record, crm_customer_tag_rel, crm_customer_share, crm_customer_protection']],
    [2, 1, 10]
)

add_heading('6.1 关键表关系', 2)
add_para('customer ←1:N→ sales_order (客户→订单)')
add_para('sales_order ←1:N→ order_item (订单→商品明细)')
add_para('portal_user → crm_customer_id → customer (会员门户用户关联CRM客户)')
add_para('crm_lead → convert_customer_id → customer (线索转客户)')
add_para('crm_opportunity → customer_id → customer (商机关联客户)')
add_para('crm_rfm_score → customer_id → customer (RFM分层关联客户)')

doc.add_page_break()

# ============================================================
# 七、API端点清单
# ============================================================
add_heading('七、API端点清单', 1)
add_para('系统共约130+个API端点，按功能域分类如下：')

add_table(
    ['功能域', '端点数', '路由前缀', '示例'],
    [['认证', '4', '/api/auth/*', 'POST /api/auth/login'],
     ['客户管理', '8', '/api/customers*', 'GET /api/customers'],
     ['线索管理', '8', '/api/leads*', 'GET /api/leads/list'],
     ['商机管理', '8', '/api/opportunities*', 'GET /api/opportunities/list'],
     ['产品管理', '6', '/api/products*', 'GET /api/products/list'],
     ['报价管理', '8', '/api/quotations*', 'POST /api/quotations/:id/convert-to-order'],
     ['订单管理', '5', '/api/orders*', 'GET /api/orders'],
     ['营销/优惠券', '10+', '/api/campaigns* /api/coupons*', 'POST /api/admin/coupons/grant'],
     ['协同办公', '12+', '多前缀', 'GET /api/schedules/list'],
     ['AI智能', '12+', '/api/ai/*', 'POST /api/ai/agent/command'],
     ['数据分析', '8+', '/api/analytics* /api/sales* /api/dashboard*', 'GET /api/rfm-scores'],
     ['CRM后台管理', '10+', '/api/admin/*', 'GET /api/admin/portal-stats'],
     ['会员门户', '40+', '/api/member-portal/*', 'POST /api/member-portal/auth/login'],
     ['学生CRM', '12+', '/api/student-profiles* /api/campus*', 'GET /api/analytics/student-dashboard']],
    [3, 2, 5, 6]
)

doc.add_page_break()

# ============================================================
# 八、UI设计体系
# ============================================================
add_heading('八、UI设计体系', 1)

add_heading('8.1 CRM后台', 2)
add_table(
    ['元素', '规格', '色值/说明'],
    [['侧边栏', '220px深蓝底', '#1E293B'],
     ['主色调', '专业蓝', '#2563EB'],
     ['强调色', '暖橙', '#F59E0B'],
     ['背景色', '浅灰蓝', '#F1F5F9'],
     ['卡片圆角', '12px', '白色/微阴影'],
     ['字体', 'Inter + PingFang SC', '14px正文']],
    [3, 3, 5]
)

add_heading('8.2 会员门户', 2)
add_table(
    ['元素', '规格', '色值/说明'],
    [['主色', '活力橙', '#FF6B35'],
     ['辅助色', '暖金', '#F7B731'],
     ['背景色', '暖灰白', '#F5F3F0'],
     ['Hero', '橙色渐变 + 装饰圆', 'linear-gradient(#FF6B35,#F7B731)'],
     ['秒杀面板', '深色玻璃效果', '#1A1A2E + 半透明'],
     ['积分Banner', '紫色渐变', 'linear-gradient(#8B5CF6,#6D28D9)'],
     ['卡片圆角', '14px', '悬停上浮4px'],
     ['按钮', '12px圆角 + 渐变色', '48px高度'],
     ['字体', 'Inter + PingFang SC', '32/20/14/12层级']],
    [3, 3, 5]
)

screenshot_hint('CRM后台 + 会员门户 左右对比截图')

doc.add_page_break()

# ============================================================
# 九、启动与部署
# ============================================================
add_heading('九、启动与部署', 1)

add_heading('9.1 环境要求', 2)
add_table(
    ['组件', '要求'],
    [['Node.js', '16+ (推荐22.22.2)'],
     ['npm', '已安装于Node.js中'],
     ['操作系统', 'Windows / macOS / Linux'],
     ['端口', '8081（后端）、3002（CRM）、3003（门户）均需可用']],
    [3, 8]
)

add_heading('9.2 一键启动', 2)
add_para('双击项目根目录的 start-crm.bat，脚本自动完成：')
add_para('1. 检测Node.js环境')
add_para('2. 清理旧端口占用（8081/3002/3003）')
add_para('3. 启动后端（node src/server.js，端口8081）')
add_para('4. 构建CRM前端 + 启动预览（端口3002）')
add_para('5. 构建会员门户 + 启动预览（端口3003）')

add_heading('9.3 手动启动（分步）', 2)
add_para('后端：')
add_para('cd crm-server && node src/server.js', bold=True)
add_para('CRM前台：')
add_para('cd original-crm/crm-frontend && npx vite build && npx vite preview --port 3002', bold=True)
add_para('门户前台：')
add_para('cd member-portal && npx vite build && npx vite preview --port 3003', bold=True)

doc.add_page_break()

# ============================================================
# 十、项目文件结构
# ============================================================
add_heading('十、项目管理文件结构', 1)
add_para('项目根目录：C:\\Users\\SMC\\WorkBuddy\\2026-05-30-23-45-28\\')
doc.add_paragraph()
add_para('├── start-crm.bat                     # 一键启动脚本')
add_para('├── CONTINE.md                        # 本报告')
add_para('├── CRM系统开发规格文档.md             # 开发规格文档')
add_para('├── crm-server/                       # 后端服务')
add_para('│   ├── src/server.js                 # ~1000行单文件后端')
add_para('│   ├── data/crm.db                   # SQLite数据库文件')
add_para('│   └── package.json')
add_para('├── original-crm/')
add_para('│   └── crm-frontend/                 # CRM管理后台')
add_para('│       └── src/')
add_para('│           ├── views/                # 36个页面视图')
add_para('│           ├── components/           # AI助手/消遣模块/布局')
add_para('│           ├── api/                  # API接口层')
add_para('│           ├── router/               # 路由配置')
add_para('│           └── styles/               # 设计系统CSS')
add_para('├── member-portal/                    # 用户会员门户')
add_para('│   └── src/')
add_para('│       ├── views/                    # 18个C端页面')
add_para('│       └── styles/                   # 高端设计系统CSS')
add_para('└── maven/                            # Maven工具（保留）')

doc.add_page_break()

# ============================================================
# 十一、数据清单
# ============================================================
add_heading('十一、数据清单', 1)
add_para('以下为系统当前实际数据量统计，截至2026年5月31日：')

add_table(
    ['序号', '模块', '数据量', '说明'],
    [['1', '客户', '12人', '含高价值/普通/沉睡各等级'],
     ['2', '产品', '28款', '8款基础 + 20款名创优品风格'],
     ['3', '订单', '56笔', '已付款/已发货/已完成等多状态'],
     ['4', '线索', '5条', '转介/展会/网站/地推/门店来源'],
     ['5', '商机', '5个', '4销售阶段，总额¥510,000'],
     ['6', '报价', '待创建', '支持6阶段审批流'],
     ['7', '营销活动', '3个', '满减/促销/新品'],
     ['8', '市场活动', '4场', '校园/会员日/新品/毕业季'],
     ['9', '优惠券', '4种', '满200-50/满100-20/无门槛5元/8折'],
     ['10', '券发放', '4条', '已发放给不同客户'],
     ['11', '会员', '4人', '金卡/银卡/普通'],
     ['12', '积分记录', '3条', '含手动调整记录'],
     ['13', 'RFM分层', '8人', '核心价值/潜力/一般/流失'],
     ['14', '日程', '7条', '客户回访日程'],
     ['15', '任务', '2条', '待办+进行中'],
     ['16', '公告', '2条', '618大促+新品上线'],
     ['17', '联系人', '5人', '关联客户1-5'],
     ['18', '服务请求', '2条', '产品咨询+退换货'],
     ['19', '反馈', '2条', '已回复1条'],
     ['20', '门户用户', '2人', '13800000000 + 13900009999'],
     ['21', '客户标签', '8种', '高频/美妆/文具/价格/新生/毕业等'],
     ['22', '工作流', '3条', '新客欢迎/沉睡唤醒/生日关怀'],
     ['23', '营销模板', '3条', '满减/半价/闪购']],
    [1, 3, 2, 6]
)

doc.add_page_break()

# ============================================================
# 十二、开发统计
# ============================================================
add_heading('十二、开发统计', 1)

add_table(
    ['维度', '数量'],
    [['CRM后台模块', '29个'],
     ['CRM前端页面', '36个'],
     ['会员门户页面', '18个'],
     ['数据库表', '75+张'],
     ['API端点', '130+个'],
     ['AI智能指令', '15种'],
     ['消遣功能', '5项'],
     ['后端代码行', '~1000行 (server.js)'],
     ['产品数据', '28款'],
     ['累计订单', '56笔'],
     ['本轮开发修复次数', '27次'],
     ['总对话轮次', '50+轮']],
    [3, 5]
)

add_para('')
add_para('系统已全面可运行。CRM后台和会员门户均已通过功能验证。')
add_para('')
add_para('— 报告结束 —', align=WD_ALIGN_PARAGRAPH.CENTER)

# ====== 保存 ======
output_path = 'c:/Users/SMC/WorkBuddy/2026-05-30-23-45-28/CRM系统开发报告_v3.0.docx'
doc.save(output_path)
print(f'报告已保存至：{output_path}')
